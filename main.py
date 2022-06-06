from docx import Document
import re
import numpy as np
class document:
#КОПИРУЕМ ПАРАМЕТРЫ ИЗ ИСХОДНОГО ДОКУМЕНТА
    @staticmethod
    def GetParaData(output_doc_name, paragraph):
        output_para = output_doc_name.add_paragraph()
        for run in paragraph.runs:
            output_run = output_para.add_run(run.text)
            output_run.bold = run.bold
            output_run.italic = run.italic
            output_run.underline = run.underline
            output_run.font.color.rgb = run.font.color.rgb
            output_run.style = run.style
            output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
            output_para.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
            output_para.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
            output_para.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
            output_para.content_type = paragraph.part.content_type
#ЗАБИРАЕМ ГОСТЫ ИЗ БАЗЫ. ПОДГОТАВЛИВАЕМ ДЛЯ ПОИСКА МАССИВ
    @staticmethod
    def BaseGost():
        import sqlite3
        conn = sqlite3.connect("sqlite.db")
        conn.row_factory = lambda cursor, row: row[0]
        c = conn.cursor()
        gost2 = c.execute('SELECT * FROM tablegost2').fetchall()
        splitgost = []
        for i in range(len(gost2)):
            g = gost2[i].split('-')
            splitgost.append(g[0])
            splitgost.append(g[1])
        buffer = np.array(splitgost)
        resbuffer = buffer.reshape(-1, 2)
        resgost2 = np.array(resbuffer).tolist()
        return resgost2
#ПОИСК УСТАРЕВШЕГО ГОДА
    def SerchGod(self,fileName):
        self.fileName=fileName
        doc = Document(fileName)

        completedText = []
        BaseGost=self.BaseGost()
        for paragraph in doc.paragraphs:
            completedText.append(paragraph.text)
        hub=[]
        findgost=[]
        god=[]
        result=[]
        lengthparag = len(completedText)
        for i in range(lengthparag):
                result4= re.findall(r"ГОСТ Р \d{5}-\d{2}", completedText[i])
                if result4!=[]:
                    hub.append(result4)

        findgost.append(sum(hub, []))
        print(findgost)
        splitbuffer=[]
        for i in range(len(findgost[0])):
            itog = findgost[0][i].split('-')
            splitbuffer.append(itog)
        print(splitbuffer)
        for i in range(len(splitbuffer)):
            for j in range(len(BaseGost)):
                if splitbuffer[i][0] == BaseGost[j][0]:
                    print('Same gost')
                    if splitbuffer[i][1] == BaseGost[j][1]:
                        print('Same god')
                    else:
                        result.append(splitbuffer[i][0])
                        god.append(BaseGost[j][1])
                        print(result)
                        print('Not same god')
        for i in range(lengthparag):
                for j in range(len(result)):
                        if completedText[i].find(result[j]) != 1:
                         char = completedText[i]
                         char = char.replace(result[j], result[j]+'-ЭТОТ ГОСТ УСТАРЕЛ,АКТУАЛЬНЫЙ ГОД:'+god[j])
                         completedText[i]=char
                         doc.paragraphs[i].text=char
                        else:
                         print('Строка ненайдена')
        return doc.paragraphs
#ПОИСК ЗАМЕНЕННОГО ГОСТА
    def SerchChange(self, fileName):
        self.fileName=fileName
        doc = Document(fileName)
        completedText = []
        BaseOldGost=['ГОСТ Р 50442-92','ГОСТ Р 8.3343.33-98']
        BaseNewGost=['ГОСТ Р 0008.003-2019','ГОСТ 0008.000-2019']
        for paragraph in doc.paragraphs:
            completedText.append(paragraph.text)
        lengthparag = len(completedText)
        for i in range(lengthparag):
            for j in range(len(BaseOldGost)):
                if completedText[i].find(BaseOldGost[j]) != 1:
                    char = completedText[i]
                    char = char.replace(BaseOldGost[j],BaseOldGost[j]+':ЗАМЕНЕН НА-'+BaseNewGost[j])
                    completedText[i] = char
                    doc.paragraphs[i].text = char
                else:
                    print('Строка ненайдена')
        return doc.paragraphs
#СОХРАНЕНИЕ ДОКУМЕНТА ЗА НОВЫМ ИМЕНЕМ
    def Save(self,Paragraph):
        length = len(Paragraph)
        Newdoc = Document()
        for i in range(length):
            self.GetParaData(Newdoc, Paragraph[i])
        for j in range(2,10):
            j=str(j)
            try:
                open('demo'+j+'.docx')
            except FileNotFoundError:
                Newdoc.save('demo'+j+'.docx')
                break
O=document()
BigO=O.SerchGod('demo.docx')
O.Save(BigO)
BigO=O.SerchChange('demo.docx')
O.Save(BigO)